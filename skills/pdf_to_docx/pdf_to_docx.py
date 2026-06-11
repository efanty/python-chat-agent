"""pdf_to_docx Skill — 使用智谱视觉大模型将 PDF 转换为 Word 文档。

工作流程:
  1. PDF → 图片 (PyMuPDF/fitz)
  2. 图片 → Markdown 结构化文本 (智谱 GLM-4V 视觉模型)
  3. Markdown → .docx 文件 (python-docx)
"""

import os
import re
import json
import base64
import requests
from pathlib import Path
from io import BytesIO

# ── 路径 ──────────────────────────────────────────────────────────────────

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_ENV_PATH = _PROJECT_ROOT / ".env"

try:
    from dotenv import load_dotenv
    load_dotenv(_ENV_PATH)
except ImportError:
    pass

# ── API 配置 ──────────────────────────────────────────────────────────────

API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
DEFAULT_MODEL = os.getenv("ZHIPU_VISION_MODEL", "glm-4v-flash")
FALLBACK_MODELS = ["glm-4v-plus", "glm-5v-turbo"]

# ── 默认提示词 ────────────────────────────────────────────────────────────

DEFAULT_DOC_PROMPT = (
    "请识别此图片中的所有文字内容，以 Markdown 格式输出。\n"
    "要求：\n"
    "1. 保留原文的标题层级（# 一级标题、## 二级标题等）\n"
    "2. 段落之间用空行分隔\n"
    "3. 表格用 Markdown 表格格式（| 列1 | 列2 |）\n"
    "4. 列表用 - 或 1. 格式\n"
    "5. 代码用 ``` 代码块包裹\n"
    "6. 保留原文的排版顺序\n"
    "7. 只输出 Markdown 内容，不要额外说明"
)


# ── 核心函数 ──────────────────────────────────────────────────────────────

def _resolve_path(file_path: str, sandbox_dir: str = None) -> Path:
    p = Path(file_path)
    if p.is_absolute():
        return p
    if sandbox_dir:
        return Path(sandbox_dir) / file_path
    return _PROJECT_ROOT / file_path


def _parse_page_range(pages_spec: str, total_pages: int) -> list[int]:
    if not pages_spec or pages_spec.lower() == "all":
        return list(range(1, total_pages + 1))
    pages = []
    for part in pages_spec.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            pages.extend(range(int(start.strip()), int(end.strip()) + 1))
        else:
            pages.append(int(part))
    return sorted(set(pages))


def _pdf_to_images(pdf_path: Path, dpi: int = 200) -> list[BytesIO]:
    """将 PDF 每页转为 PNG 图片，返回 BytesIO 列表。

    使用 PyMuPDF (fitz) 实现，无需安装 poppler。
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("需要 PyMuPDF: pip install PyMuPDF")

    doc = fitz.open(str(pdf_path))
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    bufs = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(matrix=matrix)
        img_bytes = pix.tobytes("png")
        buf = BytesIO(img_bytes)
        buf.seek(0)
        bufs.append(buf)
    doc.close()
    return bufs


def _encode_image(image_buf: BytesIO) -> str:
    b64 = base64.b64encode(image_buf.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{b64}"


def _call_vision_api(data_url: str, prompt: str, model: str, api_key: str) -> tuple[bool, str, str]:
    models_to_try = [model] + [m for m in FALLBACK_MODELS if m != model]
    last_error = ""
    for attempt_model in models_to_try:
        payload = {
            "model": attempt_model,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": data_url}},
                    {"type": "text", "text": prompt},
                ],
            }],
        }
        headers = {"Authorization": f"Bearer {api_key}"}
        try:
            resp = requests.post(API_URL, headers=headers, json=payload, timeout=120)
            if resp.status_code == 200:
                data = resp.json()
                choices = data.get("choices", [])
                if choices:
                    text = choices[0].get("message", {}).get("content", "")
                    return True, text, attempt_model
                last_error = "API 返回了空的 choices"
            else:
                last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
        except requests.exceptions.Timeout:
            last_error = "请求超时"
        except requests.exceptions.RequestException as e:
            last_error = f"请求失败: {e}"
        except Exception as e:
            last_error = f"解析错误: {e}"
    return False, last_error, model


def _markdown_to_docx(markdown_text: str, output_path: Path, title: str = "") -> str:
    """将 Markdown 文本转换为 .docx 文件。"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("需要 python-docx: pip install python-docx")

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = 'Microsoft YaHei'
        heading_font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if title:
        doc.add_heading(title, level=1)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            if in_table and table_data:
                _add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
            i += 1
            continue

        # 表格行
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # 跳过表头分隔行（如 |---|----|）
            if all(re.match(r'^[-:\s]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_data = []
            table_data.append(cells)
            i += 1
            continue
        else:
            if in_table and table_data:
                _add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # 列表项
        if re.match(r'^[\-\*]\s+', stripped):
            text = re.sub(r'^[\-\*]\s+', '', stripped)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+[\.\)]\s+', stripped):
            text = re.sub(r'^\d+[\.\)]\s+', '', stripped)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # 普通段落（支持加粗、斜体）
        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)
        i += 1

    # 处理末尾的表格
    if in_table and table_data:
        _add_table_to_doc(doc, table_data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


def _add_table_to_doc(doc, table_data: list[list[str]]):
    """将表格数据添加到文档。"""
    if not table_data:
        return
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    for r, row in enumerate(table_data):
        for c, cell_text in enumerate(row):
            if c < len(table.rows[r].cells):
                cell = table.rows[r].cells[c]
                cell.text = cell_text
                # 表头加粗
                if r == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    doc.add_paragraph()  # 表后空行


def _add_formatted_run(paragraph, text: str):
    """添加带格式（加粗、斜体）的文本到段落。"""
    # 解析 **加粗** 和 *斜体*
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ── 入口函数 ──────────────────────────────────────────────────────────────

def run(expression: str = "", action: str = "", **kwargs) -> str:
    """Top-level entry point called by SkillExecutor.

    将 PDF 文件转换为 Word 文档。

    Args:
        expression: JSON 字符串或文件路径
        action: "convert"（默认）
        **kwargs:
            file_path: PDF 文件路径（必填）
            output: 输出 Word 文件名（可选）
            pages: 页码范围（可选）
            prompt: 自定义识别提示词（可选）
            model: 视觉模型名称（可选）

    Returns:
        JSON 字符串
    """
    action = action or "convert"

    # 解析参数
    if expression and expression.strip().startswith("{"):
        try:
            args = json.loads(expression)
            file_path = args.get("file_path", kwargs.get("file_path", ""))
            output = args.get("output", kwargs.get("output", ""))
            pages = args.get("pages", kwargs.get("pages", "all"))
            prompt = args.get("prompt", kwargs.get("prompt", ""))
            model = args.get("model", kwargs.get("model", DEFAULT_MODEL))
        except json.JSONDecodeError:
            file_path = expression
            output = kwargs.get("output", "")
            pages = kwargs.get("pages", "all")
            prompt = kwargs.get("prompt", "")
            model = kwargs.get("model", DEFAULT_MODEL)
    else:
        file_path = kwargs.get("file_path", "") or expression
        output = kwargs.get("output", "")
        pages = kwargs.get("pages", "all")
        prompt = kwargs.get("prompt", "")
        model = kwargs.get("model", DEFAULT_MODEL)

    if not file_path:
        return json.dumps({
            "success": False,
            "error": "请提供 PDF 文件路径。用法: run(file_path='/path/to/file.pdf')",
        }, ensure_ascii=False)

    pdf_path = _resolve_path(file_path)
    if not pdf_path.exists():
        return json.dumps({
            "success": False,
            "error": f"PDF 文件不存在: {file_path}",
        }, ensure_ascii=False)

    if pdf_path.suffix.lower() != ".pdf":
        return json.dumps({
            "success": False,
            "error": f"不是 PDF 文件: {pdf_path.name}",
        }, ensure_ascii=False)

    api_key = os.getenv("ZHIPU_API_KEY", "")
    if not api_key:
        return json.dumps({
            "success": False,
            "error": "ZHIPU_API_KEY 未设置。请从 https://open.bigmodel.cn 获取 API Key。",
        }, ensure_ascii=False)

    sandbox_dir = kwargs.get("sandbox_dir")
    if not output:
        output = pdf_path.with_suffix(".docx").name
    output_path = _resolve_path(output, sandbox_dir)

    # ── 步骤 1: PDF 转图片 ────────────────────────────────────────────
    try:
        images = _pdf_to_images(pdf_path)
    except ImportError as e:
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"PDF 转图片失败: {e}",

        }, ensure_ascii=False)

    total_pages = len(images)
    target_pages = _parse_page_range(pages, total_pages)

    # ── 步骤 2: 逐页识别 ──────────────────────────────────────────────
    use_prompt = prompt or DEFAULT_DOC_PROMPT
    all_markdown = []
    errors = []

    for page_num in target_pages:
        if page_num < 1 or page_num > total_pages:
            errors.append(f"第 {page_num} 页不存在（共 {total_pages} 页）")
            continue

        img_buf = images[page_num - 1]
        data_url = _encode_image(img_buf)

        success, result_text, model_used = _call_vision_api(data_url, use_prompt, model, api_key)
        if not success:
            errors.append(f"第 {page_num} 页识别失败: {result_text}")
            continue

        # 添加分页标记
        if all_markdown:
            all_markdown.append(f"\n\n---\n\n")
        all_markdown.append(f"## 第 {page_num} 页\n\n{result_text.strip()}")

    if not all_markdown:
        return json.dumps({
            "success": False,
            "error": "未能从 PDF 中识别出任何内容",
            "details": errors,
            "total_pages": total_pages,
        }, ensure_ascii=False)

    full_markdown = "\n".join(all_markdown)

    # ── 步骤 3: 重建 Word ─────────────────────────────────────────────
    try:
        docx_path = _markdown_to_docx(full_markdown, output_path, title=pdf_path.stem)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"创建 Word 文档失败: {e}",
        }, ensure_ascii=False)

    result = {
        "success": True,
        "output": docx_path,
        "total_pages": total_pages,
        "processed_pages": len(target_pages),
        "model_used": model,
    }
    if errors:
        result["warnings"] = errors

    return json.dumps(result, ensure_ascii=False)


# ── Action aliases ────────────────────────────────────────────────────────

convert = run
