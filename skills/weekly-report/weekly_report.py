"""weekly_report Skill — 周报生成器。

根据本周待办完成情况、模具项目进度、会议纪要等数据，
自动汇总生成周报 Word 文档。
"""

import json
import os
import sqlite3
from datetime import datetime, timedelta, date
from pathlib import Path

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_SKILL_DIR = Path(__file__).resolve().parent
_MOLD_DB = _PROJECT_ROOT / "app" / "plugins" / "mold" / "mold_management.db"
_SANDBOX_DIR = _PROJECT_ROOT / "sandbox"


def _get_week_range(week_start: str = "", week_end: str = "") -> tuple:
    """获取周报的日期范围。"""
    today = date.today()

    if week_start:
        try:
            start = datetime.strptime(week_start, "%Y-%m-%d").date()
        except ValueError:
            start = today - timedelta(days=today.weekday())
    else:
        start = today - timedelta(days=today.weekday())

    if week_end:
        try:
            end = datetime.strptime(week_end, "%Y-%m-%d").date()
        except ValueError:
            end = start + timedelta(days=6)
    else:
        end = start + timedelta(days=6)

    return start, end


def _get_mold_conn():
    """连接模具数据库。"""
    db_path = os.path.normpath(str(_MOLD_DB))
    if not os.path.exists(db_path):
        return None
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _get_todo_data(start: date, end: date, mold_number: str = "") -> dict:
    """获取待办数据。"""
    try:
        from flask import current_app
        from app import create_app

        app = create_app()
        with app.app_context():
            from app.extensions.init_sqlalchemy import db
            from app.models.todo import Todo

            query = Todo.query
            if mold_number:
                query = query.filter_by(mold_number=mold_number)

            # 本周创建的待办
            week_created = query.filter(
                Todo.timestamp >= datetime.combine(start, datetime.min.time()),
                Todo.timestamp <= datetime.combine(end, datetime.max.time()),
            ).all()

            # 本周完成的待办
            week_done = query.filter(
                Todo.done == True,
                Todo.timestamp >= datetime.combine(start, datetime.min.time()),
                Todo.timestamp <= datetime.combine(end, datetime.max.time()),
            ).all()

            # 所有未完成的待办
            pending = query.filter_by(done=False).all()

            return {
                "created": len(week_created),
                "done": len(week_done),
                "pending": len(pending),
                "items_created": [
                    {"body": t.body, "mold_number": t.mold_number, "due_date": str(t.due_date) if t.due_date else ""}
                    for t in week_created[:20]
                ],
                "items_done": [
                    {"body": t.body, "mold_number": t.mold_number}
                    for t in week_done[:20]
                ],
            }
    except Exception as e:
        return {"created": 0, "done": 0, "pending": 0, "items_created": [], "items_done": [], "error": str(e)}


def _get_mold_data(start: date, end: date, mold_number: str = "") -> list:
    """获取模具项目数据。"""
    conn = _get_mold_conn()
    if not conn:
        return []

    try:
        projects = []
        if mold_number:
            rows = conn.execute(
                "SELECT * FROM project_info WHERE mold_number=?", (mold_number,)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM project_info ORDER BY kickoff_date DESC LIMIT 20"
            ).fetchall()

        for row in rows:
            row = dict(row)
            mn = row["mold_number"]

            # 本周试模
            trials = conn.execute(
                "SELECT * FROM trial_records WHERE mold_number=? AND trial_date >= ? AND trial_date <= ? ORDER BY trial_date",
                (mn, start.isoformat(), end.isoformat()),
            ).fetchall()

            # 本周问题
            issues = conn.execute(
                "SELECT * FROM mold_issue_records WHERE mold_number=? AND created_at >= ? AND created_at <= ? ORDER BY created_at",
                (mn, start.isoformat(), end.isoformat()),
            ).fetchall()

            # 本周改模
            modifications = conn.execute(
                "SELECT * FROM modification_records WHERE mold_number=? AND created_at >= ? AND created_at <= ? ORDER BY created_at",
                (mn, start.isoformat(), end.isoformat()),
            ).fetchall()

            projects.append({
                "mold_number": mn,
                "customer_name": row.get("customer_name", ""),
                "product_name": row.get("product_name", ""),
                "trials": [dict(t) for t in trials],
                "issues": [dict(i) for i in issues],
                "modifications": [dict(m) for m in modifications],
            })

        return projects
    finally:
        conn.close()


def _get_meeting_data(start: date, end: date) -> list:
    """获取本周会议纪要。"""
    try:
        from flask import current_app
        from app import create_app

        app = create_app()
        with app.app_context():
            from app.extensions.init_sqlalchemy import db
            from sqlalchemy import Column, Integer, String, Text, Date, DateTime

            class MeetingMinutes(db.Model):
                __tablename__ = "meeting_minutes"
                id = Column(Integer, primary_key=True)
                title = Column(String(256))
                summary = Column(String(512))
                meeting_date = Column(Date)
                created_at = Column(DateTime)

            records = MeetingMinutes.query.filter(
                MeetingMinutes.meeting_date >= start,
                MeetingMinutes.meeting_date <= end,
            ).order_by(MeetingMinutes.meeting_date.desc()).all()

            return [
                {
                    "id": r.id,
                    "title": r.title,
                    "summary": r.summary,
                    "meeting_date": str(r.meeting_date) if r.meeting_date else "",
                }
                for r in records
            ]
    except Exception:
        return []


def _generate_docx(title: str, start: date, end: date, todo_data: dict,
                   mold_projects: list, meetings: list) -> str:
    """生成周报 Word 文档。"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise RuntimeError("需要 python-docx 库: pip install python-docx")

    doc = Document()

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期信息
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"报告周期: {start.isoformat()} ~ {end.isoformat()}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    gen_time = doc.add_paragraph()
    gen_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_time.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()  # 空行

    # ── 一、本周工作概览 ─────────────────────────────────────────────
    doc.add_heading("一、本周工作概览", level=1)

    overview = doc.add_paragraph()
    overview.add_run(f"• 本周新增待办: {todo_data['created']} 项\n")
    overview.add_run(f"• 本周完成待办: {todo_data['done']} 项\n")
    overview.add_run(f"• 当前待处理: {todo_data['pending']} 项\n")
    overview.add_run(f"• 在产模具项目: {len(mold_projects)} 个\n")
    overview.add_run(f"• 本周会议: {len(meetings)} 次")

    # ── 二、模具项目进度 ─────────────────────────────────────────────
    doc.add_heading("二、模具项目进度", level=1)

    if not mold_projects:
        doc.add_paragraph("（本周无模具项目数据）")
    else:
        for proj in mold_projects:
            doc.add_heading(f"{proj['mold_number']} - {proj['customer_name']} / {proj['product_name']}", level=2)

            # 试模
            if proj["trials"]:
                p = doc.add_paragraph()
                p.add_run("试模记录:").bold = True
                for t in proj["trials"]:
                    doc.add_paragraph(
                        f"  • {t.get('trial_date', '-')} - 第{t.get('trial_count', '-')}次 - "
                        f"结果: {t.get('result', '-')}",
                        style="List Bullet",
                    )

            # 问题
            if proj["issues"]:
                p = doc.add_paragraph()
                p.add_run("新发现问题:").bold = True
                for i in proj["issues"]:
                    doc.add_paragraph(
                        f"  • {i.get('issue_description', '-')[:60]}",
                        style="List Bullet",
                    )

            # 改模
            if proj["modifications"]:
                p = doc.add_paragraph()
                p.add_run("改模记录:").bold = True
                for m in proj["modifications"]:
                    doc.add_paragraph(
                        f"  • {m.get('modification_content', '-')[:50]}",
                        style="List Bullet",
                    )

            if not proj["trials"] and not proj["issues"] and not proj["modifications"]:
                doc.add_paragraph("（本周无更新）")

    # ── 三、待办事项 ─────────────────────────────────────────────────
    doc.add_heading("三、待办事项", level=1)

    doc.add_heading("本周新增", level=2)
    if todo_data["items_created"]:
        for item in todo_data["items_created"]:
            due = f"（截止: {item['due_date']}）" if item["due_date"] else ""
            doc.add_paragraph(f"• {item['body']} {due}", style="List Bullet")
    else:
        doc.add_paragraph("（本周无新增待办）")

    doc.add_heading("本周完成", level=2)
    if todo_data["items_done"]:
        for item in todo_data["items_done"]:
            doc.add_paragraph(f"• {item['body']}", style="List Bullet")
    else:
        doc.add_paragraph("（本周无完成待办）")

    # ── 四、会议纪要 ─────────────────────────────────────────────────
    doc.add_heading("四、会议纪要", level=1)

    if meetings:
        for m in meetings:
            doc.add_heading(f"{m['meeting_date']} - {m['title']}", level=2)
            if m["summary"]:
                doc.add_paragraph(m["summary"])
    else:
        doc.add_paragraph("（本周无会议记录）")

    # ── 五、下周计划 ─────────────────────────────────────────────────
    doc.add_heading("五、下周计划", level=1)
    doc.add_paragraph("（请编辑补充下周工作计划）")

    # 保存
    _SANDBOX_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"weekly_report_{start.isoformat()}_{end.isoformat()}.docx"
    filepath = _SANDBOX_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


def _save_report(title: str, start: date, end: date, filepath: str, summary: str) -> dict:
    """保存周报记录到数据库。"""
    try:
        from flask import current_app
        from app import create_app

        app = create_app()
        with app.app_context():
            from app.extensions.init_sqlalchemy import db
            from sqlalchemy import Column, Integer, String, Text, Date, DateTime

            class WeeklyReport(db.Model):
                __tablename__ = "weekly_reports"
                id = Column(Integer, primary_key=True)
                title = Column(String(256), nullable=False)
                week_start = Column(Date, nullable=False)
                week_end = Column(Date, nullable=False)
                file_path = Column(String(512), nullable=False)
                summary = Column(String(512), default="")
                created_at = Column(db.DateTime, default=lambda: datetime.now())

            db.create_all()
            record = WeeklyReport(
                title=title,
                week_start=start,
                week_end=end,
                file_path=filepath,
                summary=summary[:500],
            )
            db.session.add(record)
            db.session.commit()
            return {
                "report_id": record.id,
                "title": title,
                "file": filepath,
                "summary": summary,
            }
    except Exception as e:
        return {
            "report_id": None,
            "title": title,
            "file": filepath,
            "summary": summary,
            "db_warning": str(e),
        }


def _list_reports(keyword: str = "", limit: int = 20) -> list:
    """查询周报列表。"""
    try:
        from flask import current_app
        from app import create_app

        app = create_app()
        with app.app_context():
            from app.extensions.init_sqlalchemy import db
            from sqlalchemy import Column, Integer, String, Text, Date, DateTime

            class WeeklyReport(db.Model):
                __tablename__ = "weekly_reports"
                id = Column(Integer, primary_key=True)
                title = Column(String(256))
                week_start = Column(Date)
                week_end = Column(Date)
                file_path = Column(String(512))
                summary = Column(String(512))
                created_at = Column(db.DateTime)

            query = WeeklyReport.query
            if keyword:
                kw = f"%{keyword}%"
                query = query.filter(
                    db.or_(
                        WeeklyReport.title.like(kw),
                        WeeklyReport.summary.like(kw),
                    )
                )
            records = query.order_by(WeeklyReport.created_at.desc()).limit(min(limit, 100)).all()

            return [
                {
                    "id": r.id,
                    "title": r.title,
                    "week_start": str(r.week_start) if r.week_start else "",
                    "week_end": str(r.week_end) if r.week_end else "",
                    "file_path": r.file_path,
                    "summary": r.summary,
                    "created_at": str(r.created_at) if r.created_at else "",
                }
                for r in records
            ]
    except Exception:
        return []


def _get_report(report_id: int) -> dict:
    """获取周报详情。"""
    try:
        from flask import current_app
        from app import create_app

        app = create_app()
        with app.app_context():
            from app.extensions.init_sqlalchemy import db
            from sqlalchemy import Column, Integer, String, Text, Date, DateTime

            class WeeklyReport(db.Model):
                __tablename__ = "weekly_reports"
                id = Column(Integer, primary_key=True)
                title = Column(String(256))
                week_start = Column(Date)
                week_end = Column(Date)
                file_path = Column(String(512))
                summary = Column(String(512))
                created_at = Column(db.DateTime)

            record = WeeklyReport.query.get(report_id)
            if not record:
                return {"error": f"未找到 ID 为 {report_id} 的周报"}

            return {
                "id": record.id,
                "title": record.title,
                "week_start": str(record.week_start) if record.week_start else "",
                "week_end": str(record.week_end) if record.week_end else "",
                "file_path": record.file_path,
                "summary": record.summary,
                "created_at": str(record.created_at) if record.created_at else "",
            }
    except Exception as e:
        return {"error": str(e)}


# ── 入口函数 ──────────────────────────────────────────────────────────────

def run(expression: str = "", action: str = "", **kwargs) -> str:
    """周报生成器。

    Args:
        expression: JSON 字符串
        action: "generate" / "list" / "get"
        **kwargs: 见 SKILL.md 参数说明

    Returns:
        JSON 字符串
    """
    # 解析参数
    params = {}
    if expression and expression.strip().startswith("{"):
        try:
            params = json.loads(expression)
        except json.JSONDecodeError:
            pass

    action = params.get("action") or action or "generate"
    mold_number = params.get("mold_number") or kwargs.get("mold_number", "")
    week_start = params.get("week_start") or kwargs.get("week_start", "")
    week_end = params.get("week_end") or kwargs.get("week_end", "")
    report_id = params.get("report_id") or kwargs.get("report_id", 0)
    keyword = params.get("keyword") or kwargs.get("keyword", "")
    limit = params.get("limit") or kwargs.get("limit", 20)

    # ── list: 查询列表 ───────────────────────────────────────────────
    if action == "list":
        try:
            records = _list_reports(keyword=keyword, limit=int(limit))
            return json.dumps({
                "success": True,
                "records": records,
                "total": len(records),
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({
                "success": False,
                "error": f"查询失败: {e}",
            }, ensure_ascii=False)

    # ── get: 查看详情 ────────────────────────────────────────────────
    if action == "get":
        try:
            report_id = int(report_id)
        except (ValueError, TypeError):
            return json.dumps({
                "success": False,
                "error": "report_id 必须是整数",
            }, ensure_ascii=False)
        record = _get_report(report_id)
        if "error" in record:
            return json.dumps({
                "success": False,
                "error": record["error"],
            }, ensure_ascii=False)
        return json.dumps({
            "success": True,
            "record": record,
        }, ensure_ascii=False)

    # ── generate: 生成周报 ───────────────────────────────────────────
    try:
        start, end = _get_week_range(week_start, week_end)
        title = f"周报 {start.isoformat()} ~ {end.isoformat()}"

        # 获取数据
        todo_data = _get_todo_data(start, end, mold_number)
        mold_projects = _get_mold_data(start, end, mold_number)
        meetings = _get_meeting_data(start, end)

        # 生成摘要
        summary_parts = []
        if todo_data["done"] > 0:
            summary_parts.append(f"本周完成待办 {todo_data['done']} 项")
        if todo_data["created"] > 0:
            summary_parts.append(f"新增 {todo_data['created']} 项")
        if mold_projects:
            active = [p["mold_number"] for p in mold_projects if p["trials"] or p["issues"]]
            if active:
                summary_parts.append(f"{', '.join(active)} 有更新")
        if meetings:
            summary_parts.append(f"会议 {len(meetings)} 次")
        summary = "；".join(summary_parts) if summary_parts else "本周无数据更新"

        # 生成 Word
        filepath = _generate_docx(title, start, end, todo_data, mold_projects, meetings)

        # 保存记录
        saved = _save_report(title, start, end, filepath, summary)

        return json.dumps({
            "success": True,
            "report_id": saved["report_id"],
            "title": title,
            "file": filepath,
            "summary": summary,
            "stats": {
                "todos_created": todo_data["created"],
                "todos_done": todo_data["done"],
                "todos_pending": todo_data["pending"],
                "projects": len(mold_projects),
                "meetings": len(meetings),
            },
        }, ensure_ascii=False)

    except RuntimeError as e:
        return json.dumps({
            "success": False,
            "error": str(e),
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"生成周报失败: {e}",
        }, ensure_ascii=False)
